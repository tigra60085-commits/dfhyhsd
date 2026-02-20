"""Pharma-compare: detailed structured drug comparison handler.

Implements the pharma-compare skill:
  • Multi-step conversation to collect drugs, context, focus, audience
  • Generates a structured Telegram comparison (≤1800 chars)
  • Generates a formatted .docx analytical document
"""

import io
from datetime import date

from telegram import Update, InputFile
from telegram.ext import ContextTypes

from states import (
    PHARMA_COMPARE_INPUT, PHARMA_COMPARE_CONTEXT,
    PHARMA_COMPARE_FOCUS, PHARMA_COMPARE_AUDIENCE, MAIN_MENU,
)
from keyboards.menus import (
    pharma_compare_focus_keyboard, pharma_compare_audience_keyboard,
    pharma_compare_result_keyboard, main_menu_keyboard, back_keyboard,
)
from data.drugs import get_drug_by_name, search_drugs


# ─── Conversation entry ───────────────────────────────────────────────────────

async def start_pharma_compare(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text(
        "🔬 *Детальный сравнительный анализ*\n\n"
        "Введите МНН препаратов через запятую (2–4 препарата):\n"
        "_Например: Флуоксетин, Эсциталопрам, Пароксетин_",
        parse_mode="Markdown",
        reply_markup=back_keyboard("back:main"),
    )
    return PHARMA_COMPARE_INPUT


async def pharma_compare_drugs_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    text = update.message.text.strip()
    drugs = [d.strip() for d in text.replace("\n", ",").split(",") if d.strip()]

    if len(drugs) < 2:
        await update.message.reply_text(
            "⚠️ Укажите *минимум 2 препарата* через запятую.",
            parse_mode="Markdown",
        )
        return PHARMA_COMPARE_INPUT

    context.user_data["pc_drugs"] = drugs[:4]
    await update.message.reply_text(
        f"Препараты: *{' / '.join(drugs[:4])}*\n\n"
        "Введите клинический контекст (нозология, тип пациента, цель):\n"
        "_Или_ /skip _для общего сравнения_",
        parse_mode="Markdown",
        reply_markup=back_keyboard("back:main"),
    )
    return PHARMA_COMPARE_CONTEXT


async def pharma_compare_context_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    text = update.message.text.strip()
    context.user_data["pc_context"] = "" if text.lower() in ("/skip", "skip") else text
    await update.message.reply_text(
        "Выберите *приоритетный фокус* сравнения:",
        parse_mode="Markdown",
        reply_markup=pharma_compare_focus_keyboard(),
    )
    return PHARMA_COMPARE_FOCUS


async def pharma_compare_focus_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    data = query.data

    if data == "back:main":
        await query.message.reply_text("Главное меню:", reply_markup=main_menu_keyboard())
        await query.message.delete()
        return MAIN_MENU

    if data.startswith("pcfocus:"):
        context.user_data["pc_focus"] = data[len("pcfocus:"):]
        await query.edit_message_text(
            "Выберите *аудиторию*:",
            parse_mode="Markdown",
            reply_markup=pharma_compare_audience_keyboard(),
        )
        return PHARMA_COMPARE_AUDIENCE

    return PHARMA_COMPARE_FOCUS


async def pharma_compare_audience_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    data = query.data

    if data == "back:main":
        await query.message.reply_text("Главное меню:", reply_markup=main_menu_keyboard())
        await query.message.delete()
        return MAIN_MENU

    if data == "pc:again":
        await query.edit_message_text(
            "Введите МНН препаратов через запятую:",
        )
        return PHARMA_COMPARE_INPUT

    if not data.startswith("pcaud:"):
        return PHARMA_COMPARE_AUDIENCE

    context.user_data["pc_audience"] = data[len("pcaud:"):]
    drugs = context.user_data.get("pc_drugs", [])
    ctx = context.user_data.get("pc_context", "")
    focus = context.user_data.get("pc_focus", "Общий обзор")
    audience = context.user_data.get("pc_audience", "resident")

    await query.edit_message_text("⏳ Генерирую сравнительный анализ...")

    records = _resolve_drugs(drugs)

    # Telegram version
    tg_text = _build_telegram_text(records, ctx, focus)
    # Split if needed (Telegram limit 4096)
    for i in range(0, len(tg_text), 4000):
        await query.message.reply_text(tg_text[i:i + 4000], parse_mode="Markdown")

    # Full docx
    try:
        docx_buf = _build_docx(records, ctx, focus, audience)
        slug = "_vs_".join(r["name"][:8].replace(" ", "") for r in records[:3])
        filename = f"compare_{slug}_v1.docx"
        await query.message.reply_document(
            document=InputFile(docx_buf, filename=filename),
            caption="📄 Полный аналитический документ (.docx)",
        )
    except ImportError:
        await query.message.reply_text(
            "⚠️ python-docx не установлен. Установите: `pip install python-docx`",
            parse_mode="Markdown",
        )
    except Exception as exc:
        await query.message.reply_text(f"⚠️ Ошибка генерации .docx: {exc}")

    await query.message.reply_text(
        "Анализ готов.",
        reply_markup=pharma_compare_result_keyboard(),
    )
    return PHARMA_COMPARE_AUDIENCE  # stay in state for "again" callback


# ─── Data helpers ─────────────────────────────────────────────────────────────

def _resolve_drugs(drug_names: list) -> list:
    records = []
    for name in drug_names:
        d = get_drug_by_name(name)
        if not d:
            results = search_drugs(name)
            d = results[0] if results else None
        if d:
            records.append({**d, "found": True, "query_name": name})
        else:
            records.append({
                "name": name, "class": "—", "mechanism": "Данные не найдены",
                "indications": [], "side_effects": [], "interactions": [],
                "dosage": "Уточнить по инструкции", "found": False, "query_name": name,
            })
    return records


_CLASS_ONSET = {
    "SSRI": "Частичный ответ 2–4 нед, полный 6–8 нед",
    "SNRI": "Частичный ответ 2–4 нед, полный 6–8 нед",
    "TCA": "2–4 нед при депрессии, 1–2 нед при нейропатии",
    "MAOI": "2–6 нед",
    "Атипичные антидепрессанты": "1–4 нед",
    "Типичные антипсихотики": "Седация — часы; антипсихотич. эффект — 2–6 нед",
    "Атипичные антипсихотики": "Седация — дни; антипсихотич. эффект — 2–6 нед",
    "Стабилизаторы настроения": "Острая мания — 5–7 дней; профилактика — 2–4 нед",
    "Бензодиазепины": "Минуты–часы",
    "Z-препараты": "30–60 мин",
    "Стимуляторы": "30–60 мин",
}

_CLASS_WITHDRAWAL = {
    "SSRI": "⚠️ Да (FINISH-симптомы), особенно при пароксетине",
    "SNRI": "⚠️ Выраженный, особенно при венлафаксине",
    "TCA": "⚠️ Умеренный",
    "MAOI": "⚠️ Есть",
    "Атипичные антидепрессанты": "Зависит от препарата",
    "Бензодиазепины": "❌ Выраженный, риск судорог",
    "Z-препараты": "⚠️ Умеренный",
    "Стабилизаторы настроения": "⚠️ Рецидив заболевания",
    "Атипичные антипсихотики": "✅ Минимальный",
    "Типичные антипсихотики": "⚠️ Холинолитическая отдача",
    "Стимуляторы": "⚠️ Слабость, сонливость",
}

_CLASS_SPECIAL = {
    "SSRI": "Пожилые: ✅ в целом безопасны; Берем.: ✅ с осторожн. (сертралин предпочтителен)",
    "SNRI": "Пожилые: ⚠️ контроль АД; Берем.: ✅ с осторожн.",
    "TCA": "Пожилые: ❌ список Бирса; Берем.: ⚠️ ограниченные данные",
    "MAOI": "Пожилые: ❌ ортостатика; Берем.: ❌ избегать",
    "Бензодиазепины": "Пожилые: ❌ список Бирса (падения); Берем.: ❌ категория D",
    "Z-препараты": "Пожилые: ⚠️ риск падений; Берем.: ⚠️ ограниченные данные",
    "Атипичные антипсихотики": "Пожилые: ⚠️ ↑смертность при деменции (black box); Берем.: ⚠️",
    "Типичные антипсихотики": "Пожилые: ❌ высокий ЭПС; Берем.: ⚠️",
    "Стабилизаторы настроения": "Берем.: Вальпроат ❌, Литий ⚠️, Ламотриджин ✅",
    "Стимуляторы": "Берем.: ❌ избегать; Пожилые: ⚠️ кардиориск",
}


def _get_cell(drug: dict, criterion: str) -> str:
    drug_class = drug.get("class", "")
    m = {
        "Класс": drug_class or "—",
        "Механизм": (drug.get("mechanism") or "—")[:130],
        "Показания": ", ".join((drug.get("indications") or [])[:3]) or "—",
        "Начало действия": _CLASS_ONSET.get(drug_class, "Индивидуально"),
        "Дозирование": drug.get("dosage") or "—",
        "Побочные эффекты": ", ".join((drug.get("side_effects") or [])[:3]) or "—",
        "Взаимодействия": ", ".join((drug.get("interactions") or [])[:2]) or "—",
        "Особые группы": _CLASS_SPECIAL.get(drug_class, "Уточнить"),
        "Синдром отмены": _CLASS_WITHDRAWAL.get(drug_class, "Уточнить"),
    }
    return m.get(criterion, "—")


# ─── Known traps (drug-pair specific) ────────────────────────────────────────

_KNOWN_TRAPS: dict[frozenset, list[str]] = {
    frozenset({"Флуоксетин", "Пароксетин"}): [
        "Пароксетин при раке молочной железы на тамоксифене — снижает активацию тамоксифена до эндоксифена через CYP2D6. Предпочтителен сертралин или эсциталопрам.",
    ],
    frozenset({"Вальпроат", "Ламотриджин"}): [
        "Вальпроат удваивает уровень ламотриджина → синдром Стивенса–Джонсона при стандартной скорости титрации. При совместном применении — вдвое медленнее наращивать дозу ламотриджина.",
    ],
    frozenset({"Оланзапин", "Арипипразол"}): [
        "При ожирении/МС — оланзапин ухудшает метаболический статус. Арипипразол метаболически нейтрален. Но: при акатизии арипипразол сам может её вызвать.",
    ],
    frozenset({"Литий", "НПВС"}): [
        "НПВС снижают почечный клиренс лития → токсичность. Пациентам на литии рекомендовать парацетамол вместо НПВС.",
    ],
    frozenset({"Клозапин", "Рисперидон"}): [
        "Рисперидон вызывает наиболее выраженную гиперпролактинемию среди атипичных; клозапин — минимальную. При сексуальных жалобах или остеопорозе у женщин — учитывать.",
    ],
}


def _generate_traps(records: list, focus: str) -> list[str]:
    names_set = frozenset(r["name"] for r in records)
    for key, traps in _KNOWN_TRAPS.items():
        if key.issubset(names_set):
            return traps

    traps = []
    classes = [r.get("class", "") for r in records]
    if "SSRI" in classes and "TCA" in classes:
        traps.append("ТЦА при передозировке кардиотоксичны — не выбирать при суицидальном риске. SSRI значительно безопаснее.")
    if "Бензодиазепины" in classes and any("антипсихотик" in c.lower() for c in classes):
        traps.append("Комбинация бензодиазепинов с клозапином (особенно парентерально) — риск остановки дыхания.")
    if not traps:
        traps.append(
            "Не переносить результаты популяционных РКИ напрямую на конкретного пациента — "
            "индивидуальные факторы (метаболизм, коморбидность, принимаемые препараты) критичны."
        )
    return traps


def _generate_pearl(records: list, focus: str) -> str:
    classes = {r.get("class", "") for r in records}
    names = {r["name"] for r in records}
    if "SSRI" in classes and "Флуоксетин" in names:
        return (
            "Флуоксетин — единственный SSRI с T½ 1–4 дня + активный метаболит норфлуоксетин "
            "T½ 7–15 дней → минимальный синдром отмены. Но при переходе на MAOI — 5 недель отмывания."
        )
    if "Стабилизаторы настроения" in classes and "Ламотриджин" in names:
        return (
            "Ламотриджин — лучший стабилизатор для профилактики депрессивных фаз при БАР, "
            "но неэффективен при острой мании. Медленная титрация — не опциональна."
        )
    if "Атипичные антипсихотики" in classes:
        return (
            "Метаболический риск атипичных антипсихотиков (от наибольшего к меньшему): "
            "Клозапин ≈ Оланзапин > Кветиапин > Рисперидон > Арипипразол ≈ Зипразидон ≈ Луразидон."
        )
    if focus == "Безопасность у пожилых":
        return "У пожилых — начинать с ½ стандартной дозы (start low, go slow). Избегать ТЦА и бензодиазепинов (список Бирса)."
    if focus == "Беременность и лактация":
        return "Для грудного вскармливания: сертралин и пароксетин имеют наименьший RID (<2%) среди антидепрессантов."
    return (
        "'Лучший препарат' не существует абстрактно — он лучший для конкретного пациента "
        "с конкретными характеристиками. Всегда начинать с профиля пациента, а не с препарата."
    )


def _generate_scenarios(records: list, ctx: str, focus: str) -> list[dict]:
    scenarios = []
    if len(records) < 2:
        return scenarios

    r_first = records[0]
    r_second = records[1]

    # Standard adult
    scenarios.append({
        "title": "Стандартный взрослый пациент",
        "choice": r_first["name"],
        "rationale": (
            f"{r_first['name']} ({r_first.get('class', '?')}): "
            f"{(r_first.get('indications') or ['—'])[0]}. "
            f"Хорошо изучен, предсказуемый профиль переносимости."
        ),
        "alternative": r_second["name"],
    })

    # Elderly
    elderly_ok = [r for r in records if r.get("class") not in ("TCA", "MAOI", "Бензодиазепины")]
    scenarios.append({
        "title": "Пожилой пациент (≥65 лет) / соматическая коморбидность",
        "choice": elderly_ok[0]["name"] if elderly_ok else "Уточнить",
        "rationale": "Избегать ТЦА и бензодиазепинов (список Бирса, риск падений, когнитивные нарушения). Начинать с ½ стандартной дозы.",
        "alternative": "Любой SSRI при отсутствии противопоказаний",
    })

    # Pregnancy
    scenarios.append({
        "title": "Беременность / репродуктивный возраст",
        "choice": "Сертралин (антидепрессант) / Ламотриджин (стабилизатор)",
        "rationale": "Наиболее изученные при беременности. Вальпроат — абсолютно противопоказан (тератогенность). Решение — с акушером-гинекологом.",
        "alternative": "Флуоксетин (при невозможности сертралина)",
    })

    # Treatment-resistant
    scenarios.append({
        "title": "Резистентный пациент (≥2 неэффективных курса)",
        "choice": "Смена класса или аугментация",
        "rationale": (
            "При шизофрении — клозапин. "
            "При депрессии — аугментация литием, атипичным антипсихотиком или ЭСТ. "
            "Наращивание дозы неэффективного препарата нецелесообразно."
        ),
        "alternative": "ЭСТ при тяжёлой резистентной депрессии",
    })

    return scenarios


def _get_evidence_base(records: list) -> list[str]:
    evidence = []
    classes = {r.get("class", "") for r in records}
    if "SSRI" in classes or "SNRI" in classes:
        evidence.append("[Cipriani et al., Lancet 2018] — 522 РКИ, 116 тыс. пациентов: все современные антидепрессанты эффективнее плацебо; сертралин — оптимальный баланс эффективности и переносимости. Уровень A")
        evidence.append("[CANMAT 2023] — SSRI и SNRI — препараты первой линии при депрессии. Уровень A")
    if "Атипичные антипсихотики" in classes or "Типичные антипсихотики" in classes:
        evidence.append("[Huhn et al., Lancet 2019] — Сравнение 32 антипсихотиков: различия в переносимости значительнее, чем в эффективности. Уровень A")
        evidence.append("[Leucht et al., Lancet 2013] — Клозапин наиболее эффективен при рефрактерной шизофрении. Уровень A")
    if "Стабилизаторы настроения" in classes:
        evidence.append("[BAP guidelines, J Psychopharmacol 2016] — Литий — gold standard профилактики при БАР. Уровень A")
        evidence.append("[Geddes et al., Lancet 2010] — Ламотриджин: профилактика депрессивных фаз БАР (NNT=5), слабо при маниакальных. Уровень A")
    if "Бензодиазепины" in classes:
        evidence.append("[Lader, Drugs 2011] — Зависимость к бензодиазепинам после 4–6 нед. Рекомендуются курсы ≤2–4 нед. Уровень B")
    evidence.append("[Клинические рекомендации МЗ РФ] — Российские стандарты по нозологии. Уточнить актуальную версию на сайте МЗ РФ.")
    return evidence


# ─── Telegram text builder ────────────────────────────────────────────────────

def _build_telegram_text(records: list, ctx: str, focus: str) -> str:
    names = " vs ".join(r["name"] for r in records)
    lines = [f"⚖️ *{names}*", "Сравнительный разбор", ""]

    if ctx:
        lines.append(f"📌 *Контекст:* {ctx}")
        lines.append("")

    lines.append(f"🔎 *Фокус:* {focus}")
    lines.append("")

    # Mechanisms
    lines.append("🔑 *Механизмы и классы:*")
    for r in records:
        mech = r.get("mechanism", "—")
        short = mech[:100] + "…" if len(mech) > 100 else mech
        lines.append(f"• *{r['name']}* ({r.get('class', '—')}): {short}")
    lines.append("")

    # Preferred scenarios per drug
    for r in records:
        lines.append(f"✅ *{r['name']}* предпочтителен при:")
        for ind in (r.get("indications") or [])[:3]:
            lines.append(f"  • {ind}")
    lines.append("")

    # Traps
    traps = _generate_traps(records, focus)
    lines.append(f"⚠️ *Ловушка:*")
    lines.append(traps[0] if traps else "—")
    lines.append("")

    # Pearl
    pearl = _generate_pearl(records, focus)
    lines.append(f"💡 *Жемчужина:*")
    lines.append(pearl)
    lines.append("")

    # Tags
    tags = " ".join(f"#{r['name'].lower().replace(' ', '_')}" for r in records[:3])
    lines.append(f"#психофармакология #сравнение {tags}")

    text = "\n".join(lines)
    if len(text) > 1800:
        text = text[:1797] + "…"
    return text


# ─── DOCX builder ─────────────────────────────────────────────────────────────

def _build_docx(records: list, ctx: str, focus: str, audience: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Title ──
    t = doc.add_heading("СРАВНИТЕЛЬНЫЙ АНАЛИЗ", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    names = " vs ".join(r["name"] for r in records)
    h = doc.add_heading(names, 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for label, value in [
        ("Клинический контекст", ctx or "Общий"),
        ("Оси сравнения", focus),
        ("Аудитория", "Ординатор" if audience == "resident" else "Специалист"),
        ("Дата", date.today().strftime("%d.%m.%Y")),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    doc.add_page_break()

    # ── Section 1: Drug profiles ──
    doc.add_heading("Раздел 1. Профили препаратов", 1)
    for r in records:
        doc.add_heading(r["name"], 2)
        for label, key in [
            ("Класс", "class"), ("Механизм действия", "mechanism"), ("Дозировка", "dosage"),
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(r.get(key) or "—"))
        for list_label, list_key in [("Показания", "indications"), ("Побочные эффекты", "side_effects")]:
            p = doc.add_paragraph()
            p.add_run(f"{list_label}: ").bold = True
            items = r.get(list_key) or []
            p.add_run(", ".join(items[:5]) if items else "—")
        doc.add_paragraph()

    # ── Section 2: Comparison table ──
    doc.add_heading("Раздел 2. Сравнительная таблица", 1)
    criteria = [
        "Класс", "Механизм", "Показания", "Начало действия",
        "Дозирование", "Побочные эффекты", "Взаимодействия",
        "Особые группы", "Синдром отмены",
    ]
    table = doc.add_table(rows=len(criteria) + 1, cols=len(records) + 1)
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    hdr_row.cells[0].text = "Критерий"
    for i, r in enumerate(records):
        cell = hdr_row.cells[i + 1]
        cell.text = r["name"]
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for row_idx, criterion in enumerate(criteria):
        row = table.rows[row_idx + 1]
        row.cells[0].text = criterion
        for col_idx, r in enumerate(records):
            row.cells[col_idx + 1].text = _get_cell(r, criterion)

    doc.add_paragraph()

    # ── Section 3: Clinical scenarios ──
    doc.add_heading("Раздел 3. Клинические сценарии", 1)
    for i, sc in enumerate(_generate_scenarios(records, ctx, focus), 1):
        doc.add_heading(f"Сценарий {i}: {sc['title']}", 3)
        for label, key in [("Выбор", "choice"), ("Обоснование", "rationale"), ("Альтернатива", "alternative")]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(sc.get(key) or "—")
        doc.add_paragraph()

    # ── Section 4: Key differences ──
    doc.add_heading("Раздел 4. Ключевые различия", 1)
    key_diffs = [
        {
            "statement": "Разные механизмы → разные клинические ниши.",
            "rationale": "Препараты разных классов не взаимозаменяемы: эффективность зависит от патофизиологии у конкретного пациента.",
            "application": "При неэффективности одного класса — переходить на другой, а не наращивать дозу.",
        },
        {
            "statement": "Профиль переносимости определяет выбор при сопутствующих состояниях.",
            "rationale": "Антихолинергические эффекты критичны у пожилых; метаболические — при ожирении; седация — при требовательных к концентрации профессиях.",
            "application": "Всегда сопоставлять побочные эффекты с уязвимыми системами конкретного пациента.",
        },
        {
            "statement": "CYP-потенциал взаимодействий различается существенно.",
            "rationale": "Пароксетин и флуоксетин — мощные ингибиторы CYP2D6; сертралин и эсциталопрам — минимальное влияние.",
            "application": "При политерапии предпочитать препараты с минимальным CYP-влиянием.",
        },
    ]
    for diff in key_diffs:
        p = doc.add_paragraph()
        p.add_run("💡 ").bold = True
        p.add_run(diff["statement"]).bold = True
        doc.add_paragraph(f"Обоснование: {diff['rationale']}")
        doc.add_paragraph(f"Применение: {diff['application']}")
        doc.add_paragraph()

    # ── Section 5: Traps ──
    doc.add_heading("Раздел 5. Подводные камни", 1)
    for trap in _generate_traps(records, focus):
        p = doc.add_paragraph()
        p.add_run("❌ ").bold = True
        p.add_run(trap)

    # ── Section 6: Evidence ──
    doc.add_heading("Раздел 6. Доказательная база", 1)
    for ev in _get_evidence_base(records):
        doc.add_paragraph(ev, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
